# api/app.py
import os
import json
from dotenv import load_dotenv

load_dotenv()

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, timedelta
import jwt
from functools import wraps
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor

app = Flask(__name__)

# Configurações
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'yukiridensuki-secret-key-2024')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Se não tiver DATABASE_URL, usa SQLite
if not app.config['SQLALCHEMY_DATABASE_URI']:
    print("❌ DATABASE_URL não encontrada! Usando SQLite para teste...")
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///curriculum.db'

db = SQLAlchemy(app)
CORS(app, 
     origins=['https://densuki.github.io', 
              'http://localhost:3000', 
              'https://studious-goggles-v6wpp65rvwwcxgg6-5000.app.github.dev',
              'https://portifolio-pj8c.onrender.com',
              '*'],  # Para teste, depois restrinja
     allow_headers=['Content-Type', 'Authorization'],
     expose_headers=['Content-Type', 'Authorization'],
     supports_credentials=True)

# ============================================
# FUNÇÕES PARA CARREGAR DADOS DOS JSONs
# ============================================
def load_json_data(filename, default_factory):
    """Carrega dados de um arquivo JSON"""
    json_path = os.path.join(os.path.dirname(__file__), '..', '..', 'docs', 'data', filename)
    
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"⚠️ Arquivo não encontrado: {json_path}")
        return None
    except json.JSONDecodeError as e:
        print(f"❌ Erro ao ler JSON: {e}")
        return None

def save_json_data(filename, data):
    """Salva dados em um arquivo JSON"""
    json_path = os.path.join(os.path.dirname(__file__), '..', '..', 'docs', 'data', filename)
    
    try:
        # Criar diretório se não existir
        os.makedirs(os.path.dirname(json_path), exist_ok=True)
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"✅ {filename} atualizado em: {json_path}")
        return True
    except Exception as e:
        print(f"⚠️ Não foi possível atualizar o JSON: {e}")
        return False

def load_curriculum_from_json():
    """Carrega os dados do curriculum.json"""
    return load_json_data('curriculum.json', None)

def load_about_from_json():
    """Carrega os dados do about.json"""
    return load_json_data('about.json', None)

def get_default_curriculum_data():
    """Retorna dados padrão para currículo"""
    return {
        "contact": {
            "name": "João Gabriel Sousa Santos",
            "location": "Fortaleza - CE",
            "phone": "(85) 9 9217-1191",
            "email": "joaogabriel4175@gmail.com",
            "linkedin": "linkedin.com/in/densuki/",
            "github": "github.com/Densuki"
        },
        "objective": "Atuar nas áreas de atendimento ao cliente, rotinas administrativas e suporte operacional, contribuindo com organização, resolução de problemas e qualidade no atendimento.",
        "summary": "Profissional com experiência em atendimento ao público e suporte em Lan House, além de atuação voluntária em ambiente escolar com apoio administrativo e organizacional. Possuo conhecimentos em informática, Pacote Office e organização documental, buscando oportunidade para desenvolver carreira na área administrativa.",
        "education": [],
        "courses": [],
        "skills": {},
        "experience": [],
        "languages": []
    }

def get_default_about_data():
    """Retorna dados padrão para about"""
    return {
        "bio": [
            "Olá! 👋 Me chamo **João Gabriel**, tenho {{age}} anos e sou natural de {{location}}.",
            "Sou um **Desenvolvedor Full Stack** e **Artista Digital** apaixonado por tecnologia, criatividade e histórias.",
            "Acredito que a tecnologia pode ser uma ferramenta poderosa para conectar pessoas e criar experiências significativas."
        ],
        "objective": "Desenvolver soluções inovadoras que unam tecnologia e arte, impactando positivamente a vida das pessoas.",
        "status": {
            "working": "Desenvolvedor Full Stack & Artista",
            "studying": "React, Node.js e Design de Interfaces"
        },
        "description": "Sou um profissional versátil que transita entre o **código** e a **arte**. Minha jornada começou no mundo do desenvolvimento web, onde descobri a paixão por criar interfaces intuitivas e funcionais.\n\nAlém da programação, sou **artista digital** e **mangaká**, explorando diferentes formas de expressão criativa. Essa combinação única me permite abordar projetos com uma visão holística, unindo técnica e estética.",
        "health": "❤️ Saudável, pratico exercícios regularmente e mantenho uma alimentação equilibrada.",
        "skills": {
            "core": ["Comunicação", "Liderança", "Trabalho em Equipe", "Resolução de Problemas", "Criatividade"],
            "technical": ["JavaScript", "Python", "HTML5", "CSS3", "React", "Node.js", "Git"],
            "creative": ["Desenho", "Ilustração Digital", "Mangá", "Escrita Criativa", "Música"],
            "languages": ["Português (Nativo)", "Inglês (Avançado)", "Espanhol (Básico)"]
        },
        "interests": ["Anime", "Manga", "Programação", "Música", "Games", "Design", "Fotografia", "Viagens"],
        "badges": ["Desenvolvedor", "Artista", "Mangaká", "Escritor", "Músico", "Gamer", "Leitor", "Viajante"],
        "history": [
            "Minha história começou em **{{birthday}}**, quando nasci em {{location}}. Desde criança, sempre fui curioso e apaixonado por histórias.",
            "Aos {{age}} anos, descobri o mundo da programação e me encantei com a possibilidade de criar coisas novas a partir do zero.",
            "Ao longo dos anos, explorei diferentes áreas, desde o **desenvolvimento web** até a **arte digital**, sempre buscando unir tecnologia e criatividade.",
            "Hoje, continuo aprendendo e evoluindo, acreditando que cada dia é uma oportunidade para criar algo incrível."
        ]
    }

# ============================================
# MODELOS DO BANCO DE DADOS
# ============================================
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    last_login = db.Column(db.DateTime)
    is_active = db.Column(db.Boolean, default=True)

class CurriculumVersion(db.Model):
    __tablename__ = 'curriculum_versions'
    id = db.Column(db.Integer, primary_key=True)
    version = db.Column(db.String(20), nullable=False)
    data = db.Column(db.JSON, nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_current = db.Column(db.Boolean, default=False)

class CurriculumEdit(db.Model):
    __tablename__ = 'curriculum_edits'
    id = db.Column(db.Integer, primary_key=True)
    field = db.Column(db.String(100), nullable=False)
    old_value = db.Column(db.Text)
    new_value = db.Column(db.Text)
    edited_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    edited_at = db.Column(db.DateTime, default=datetime.utcnow)
    version_id = db.Column(db.Integer, db.ForeignKey('curriculum_versions.id'))

# ============================================
# MODELOS PARA ABOUT
# ============================================
class AboutVersion(db.Model):
    __tablename__ = 'about_versions'
    id = db.Column(db.Integer, primary_key=True)
    version = db.Column(db.String(20), nullable=False)
    data = db.Column(db.JSON, nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_current = db.Column(db.Boolean, default=False)

class AboutEdit(db.Model):
    __tablename__ = 'about_edits'
    id = db.Column(db.Integer, primary_key=True)
    field = db.Column(db.String(100), nullable=False)
    old_value = db.Column(db.Text)
    new_value = db.Column(db.Text)
    edited_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    edited_at = db.Column(db.DateTime, default=datetime.utcnow)
    version_id = db.Column(db.Integer, db.ForeignKey('about_versions.id'))

# ============================================
# MODELO PARA PROFILE (dados do perfil)
# ============================================
class ProfileVersion(db.Model):
    __tablename__ = 'profile_versions'
    id = db.Column(db.Integer, primary_key=True)
    version = db.Column(db.String(20), nullable=False)
    data = db.Column(db.JSON, nullable=False)
    created_by = db.Column(db.Integer, db.ForeignKey('users.id'))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    is_current = db.Column(db.Boolean, default=False)

# ============================================
# DECORADOR DE AUTENTICAÇÃO
# ============================================
def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization')
        if not token:
            return jsonify({'message': 'Token é obrigatório!'}), 401
        try:
            token = token.split(' ')[1] if ' ' in token else token
            data = jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
            current_user = User.query.filter_by(id=data['user_id']).first()
            if not current_user:
                return jsonify({'message': 'Usuário não encontrado!'}), 401
        except jwt.ExpiredSignatureError:
            return jsonify({'message': 'Token expirado!'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'message': 'Token inválido!'}), 401
        return f(current_user, *args, **kwargs)
    return decorated

# ============================================
# ROTAS DE AUTENTICAÇÃO
# ============================================
@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    
    user = User.query.filter_by(username=username).first()
    if not user or user.password != password:
        return jsonify({'message': 'Credenciais inválidas!'}), 401
    
    user.last_login = datetime.utcnow()
    db.session.commit()
    
    token = jwt.encode({
        'user_id': user.id,
        'username': user.username,
        'exp': datetime.utcnow() + timedelta(hours=24)
    }, app.config['SECRET_KEY'], algorithm='HS256')
    
    return jsonify({
        'token': token,
        'user': {
            'id': user.id,
            'username': user.username,
            'email': user.email
        }
    })

@app.route('/api/auth/register', methods=['POST'])
def register():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    email = data.get('email')
    
    if User.query.filter_by(username=username).first():
        return jsonify({'message': 'Usuário já existe!'}), 400
    
    if User.query.filter_by(email=email).first():
        return jsonify({'message': 'Email já cadastrado!'}), 400
    
    user = User(username=username, password=password, email=email)
    db.session.add(user)
    db.session.commit()
    
    return jsonify({'message': 'Usuário criado com sucesso!'}), 201

@app.route('/api/auth/verify', methods=['GET'])
@token_required
def verify(current_user):
    return jsonify({
        'authenticated': True,
        'user': {
            'id': current_user.id,
            'username': current_user.username,
            'email': current_user.email
        }
    })

# ============================================
# ROTAS DO CURRÍCULO
# ============================================
@app.route('/api/curriculum', methods=['GET'])
def get_curriculum():
    # Tentar buscar do banco de dados primeiro
    version = CurriculumVersion.query.filter_by(is_current=True).first()
    
    if version:
        return jsonify(version.data)
    
    # Se não houver versão no banco, carregar do JSON
    print("📂 Nenhuma versão no banco. Carregando do curriculum.json...")
    curriculum_data = load_curriculum_from_json()
    
    if not curriculum_data:
        print("⚠️ Usando dados padrão...")
        curriculum_data = get_default_curriculum_data()
    
    # Salvar no banco como versão inicial
    version = CurriculumVersion(
        version='1.0.0',
        data=curriculum_data,
        is_current=True
    )
    db.session.add(version)
    db.session.commit()
    print("✅ Versão inicial criada no banco de dados!")
    
    return jsonify(curriculum_data)

@app.route('/api/curriculum', methods=['PUT'])
@token_required
def update_curriculum(current_user):
    data = request.json
    
    # Criar nova versão
    version = CurriculumVersion(
        version=f"1.0.{CurriculumVersion.query.count() + 1}",
        data=data,
        is_current=True,
        created_by=current_user.id
    )
    
    # Desmarcar versões anteriores como correntes
    db.session.query(CurriculumVersion).update({CurriculumVersion.is_current: False})
    
    db.session.add(version)
    db.session.commit()
    
    # Registrar edições
    old_version = CurriculumVersion.query.filter_by(is_current=False).order_by(CurriculumVersion.created_at.desc()).first()
    if old_version:
        for key in data:
            if key in old_version.data and data[key] != old_version.data[key]:
                edit = CurriculumEdit(
                    field=key,
                    old_value=str(old_version.data[key]),
                    new_value=str(data[key]),
                    edited_by=current_user.id,
                    version_id=version.id
                )
                db.session.add(edit)
        db.session.commit()
    
    # Também atualizar o arquivo JSON
    save_json_data('curriculum.json', data)
    
    return jsonify({
        'message': 'Currículo atualizado com sucesso!',
        'version': version.version
    })

# ============================================
# ROTAS DO ABOUT
# ============================================
@app.route('/api/about', methods=['GET'])
def get_about():
    """Obtém os dados do about"""
    # Tentar buscar do banco de dados primeiro
    version = AboutVersion.query.filter_by(is_current=True).first()
    
    if version:
        return jsonify(version.data)
    
    # Se não houver versão no banco, carregar do JSON
    print("📂 Nenhuma versão no banco. Carregando do about.json...")
    about_data = load_about_from_json()
    
    if not about_data:
        print("⚠️ Usando dados padrão...")
        about_data = get_default_about_data()
    
    # Salvar no banco como versão inicial
    version = AboutVersion(
        version='1.0.0',
        data=about_data,
        is_current=True
    )
    db.session.add(version)
    db.session.commit()
    print("✅ Versão inicial do About criada no banco de dados!")
    
    return jsonify(about_data)

@app.route('/api/about', methods=['PUT'])
@token_required
def update_about(current_user):
    """Atualiza os dados do about"""
    data = request.json
    
    # Criar nova versão
    version = AboutVersion(
        version=f"1.0.{AboutVersion.query.count() + 1}",
        data=data,
        is_current=True,
        created_by=current_user.id
    )
    
    # Desmarcar versões anteriores como correntes
    db.session.query(AboutVersion).update({AboutVersion.is_current: False})
    
    db.session.add(version)
    db.session.commit()
    
    # Registrar edições
    old_version = AboutVersion.query.filter_by(is_current=False).order_by(AboutVersion.created_at.desc()).first()
    if old_version:
        for key in data:
            if key in old_version.data and data[key] != old_version.data[key]:
                edit = AboutEdit(
                    field=key,
                    old_value=str(old_version.data[key]),
                    new_value=str(data[key]),
                    edited_by=current_user.id,
                    version_id=version.id
                )
                db.session.add(edit)
        db.session.commit()
    
    # Também atualizar o arquivo JSON
    save_json_data('about.json', data)
    
    return jsonify({
        'message': 'Sobre atualizado com sucesso!',
        'version': version.version
    })

@app.route('/api/about/history', methods=['GET'])
@token_required
def get_about_history(current_user):
    """Obtém o histórico de versões do about"""
    versions = AboutVersion.query.order_by(AboutVersion.created_at.desc()).limit(10).all()
    return jsonify([{
        'version': v.version,
        'created_at': v.created_at.isoformat(),
        'created_by': User.query.get(v.created_by).username if v.created_by else 'Sistema',
        'is_current': v.is_current
    } for v in versions])

# ============================================
# ROTAS DO PROFILE
# ============================================
@app.route('/api/profile', methods=['GET'])
def get_profile():
    """Obtém os dados do perfil (profile.json)"""
    # Tentar buscar do banco de dados primeiro
    version = ProfileVersion.query.filter_by(is_current=True).first()
    
    if version:
        return jsonify(version.data)
    
    # Se não houver versão no banco, carregar do JSON
    print("📂 Nenhuma versão no banco. Carregando do profile.json...")
    profile_data = load_json_data('profile.json', None)
    
    if not profile_data:
        print("⚠️ Profile.json não encontrado!")
        return jsonify({})
    
    # Salvar no banco como versão inicial
    version = ProfileVersion(
        version='1.0.0',
        data=profile_data,
        is_current=True
    )
    db.session.add(version)
    db.session.commit()
    print("✅ Versão inicial do Profile criada no banco de dados!")
    
    return jsonify(profile_data)

# ============================================
# ROTA PARA DOWNLOAD DOS ARQUIVOS
# ============================================
@app.route('/api/download/<file_type>', methods=['GET'])
def download_file(file_type):
    # Mapear extensões
    ext_map = {
        'pdf': 'pdf',
        'docx': 'docx',
        'md': 'md'
    }
    
    ext = ext_map.get(file_type, file_type)
    return jsonify({
        'url': f'https://raw.githubusercontent.com/Densuki/densuki.github.io/main/docs/assets/curriculo_pessoal.{ext}'
    })

# ============================================
# ROTA DE STATUS
# ============================================
@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'ok',
        'timestamp': datetime.utcnow().isoformat(),
        'database': 'connected' if app.config['SQLALCHEMY_DATABASE_URI'] else 'using_sqlite',
        'endpoints': {
            'curriculum': '/api/curriculum',
            'about': '/api/about',
            'profile': '/api/profile',
            'auth': '/api/auth/login, /api/auth/verify'
        }
    })

# ============================================
# ROTA PARA DOWNLOAD DO DOCX VIA PYTHON
# ============================================
@app.route('/api/curriculum/download/docx', methods=['POST'])
def download_docx():
    try:
        data = request.json
        
        # Criar documento
        doc = Document()
        
        # Configurar margens
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Título
        title = doc.add_heading('JOÃO GABRIEL SOUSA SANTOS', 0)
        title.alignment = 1  # Centralizado
        
        # Contato
        if data.get('contact'):
            contact_texts = []
            for key, value in data['contact'].items():
                if value:
                    contact_texts.append(str(value))
            if contact_texts:
                contact = doc.add_paragraph(' • '.join(contact_texts))
                contact.alignment = 1
        
        doc.add_paragraph()
        
        # Objetivo
        if data.get('objective'):
            doc.add_heading('OBJETIVO PROFISSIONAL', 2)
            doc.add_paragraph(data['objective'])
        
        # Resumo
        if data.get('summary'):
            doc.add_heading('RESUMO PROFISSIONAL', 2)
            doc.add_paragraph(data['summary'])
        
        # Formação Acadêmica
        if data.get('education'):
            doc.add_heading('FORMAÇÃO ACADÊMICA', 2)
            for edu in data['education']:
                text = f"• {edu.get('degree', '')}"
                if edu.get('institution'):
                    text += f" – {edu.get('institution')}"
                if edu.get('period'):
                    text += f" ({edu.get('period')})"
                if edu.get('status'):
                    status_labels = {
                        'concluido': 'Concluído',
                        'completo': 'Concluído',
                        'incompleto': 'Em andamento',
                        'trancado': 'Trancado',
                        'interrompido': 'Interrompido'
                    }
                    text += f" - {status_labels.get(edu['status'], edu['status'])}"
                doc.add_paragraph(text)
        
        # Cursos e Certificações
        if data.get('courses'):
            doc.add_heading('CURSOS E CERTIFICAÇÕES', 2)
            for course in data['courses']:
                text = f"• {course.get('title', '')}"
                if course.get('duration'):
                    text += f" ({course.get('duration')})"
                if course.get('institution'):
                    text += f" – {course.get('institution')}"
                if course.get('status'):
                    status_labels = {
                        'concluido': '✅ Concluído',
                        'andamento': '🔄 Em andamento',
                        'planejado': '📋 Planejado',
                        'cancelado': '❌ Cancelado',
                        'pendente': '⏳ Pendente'
                    }
                    text += f" - {status_labels.get(course['status'], course['status'])}"
                doc.add_paragraph(text)
                
                if course.get('description'):
                    # Processar markdown simples para descrição
                    desc = course['description']
                    # Remover markdown básico para texto plano
                    desc = desc.replace('**', '').replace('*', '')
                    desc = desc.replace('# ', '').replace('## ', '').replace('### ', '')
                    # Converter listas
                    lines = desc.split('\n')
                    for line in lines:
                        if line.strip().startswith('-') or line.strip().startswith('*'):
                            doc.add_paragraph(f"   {line.strip()}", style='Normal')
                        elif line.strip():
                            doc.add_paragraph(line.strip(), style='Normal')
        
        # Competências
        if data.get('skills'):
            doc.add_heading('COMPETÊNCIAS', 2)
            skills_text = []
            category_labels = {
                'core': 'Soft Skills',
                'technical': 'Hard Skills'
            }
            for category, items in data['skills'].items():
                if items:
                    label = category_labels.get(category, category.capitalize())
                    skills_text.append(f"{label}: {', '.join(items)}")
            doc.add_paragraph(' • '.join(skills_text))
        
        # Experiências Profissionais
        if data.get('experience'):
            doc.add_heading('EXPERIÊNCIAS PROFISSIONAIS', 2)
            for exp in data['experience']:
                p = doc.add_paragraph()
                p.add_run(f"Empresa: ").bold = True
                p.add_run(exp.get('company', ''))
                
                p = doc.add_paragraph()
                p.add_run(f"Cargo: ").bold = True
                p.add_run(exp.get('position', ''))
                
                if exp.get('location'):
                    p = doc.add_paragraph()
                    p.add_run(f"Localização: ").bold = True
                    p.add_run(exp.get('location', ''))
                
                p = doc.add_paragraph()
                p.add_run(f"Período: ").bold = True
                p.add_run(f"{exp.get('startDate', '')} - {exp.get('endDate', '')}")
                
                for resp in exp.get('responsibilities', []):
                    doc.add_paragraph(f"• {resp}", style='Normal')
                
                doc.add_paragraph()
        
        # Idiomas
        if data.get('languages'):
            doc.add_heading('IDIOMAS', 2)
            for lang in data['languages']:
                doc.add_paragraph(f"{lang.get('language', '')}: {lang.get('proficiency', '')}")
        
        # Salvar em memória
        file_bytes = BytesIO()
        doc.save(file_bytes)
        file_bytes.seek(0)
        
        return Response(
            file_bytes.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': 'attachment; filename=Curriculo_JoaoGabriel.docx',
                'Content-Type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }
        )
        
    except Exception as e:
        print(f"❌ Erro ao gerar DOCX: {e}")
        return jsonify({'error': str(e)}), 500

# ============================================
# INICIALIZAÇÃO
# ============================================
if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Criar usuário padrão se não existir
        if not User.query.filter_by(username='yukiridensuki').first():
            default_user = User(
                username='yukiridensuki',
                password='yukiridensuki4175',
                email='joaogabriel4175@gmail.com',
                is_active=True
            )
            db.session.add(default_user)
            db.session.commit()
            print("✅ Usuário padrão criado: yukiridensuki / yukiridensuki4175")
        
        # Carregar dados iniciais do currículo se não houver versões
        if CurriculumVersion.query.count() == 0:
            print("📂 Nenhuma versão do currículo encontrada. Carregando dados iniciais...")
            curriculum_data = load_curriculum_from_json()
            if not curriculum_data:
                curriculum_data = get_default_curriculum_data()
            
            initial_version = CurriculumVersion(
                version='1.0.0',
                data=curriculum_data,
                is_current=True
            )
            db.session.add(initial_version)
            print("✅ Dados iniciais do currículo carregados!")
        
        # Carregar dados iniciais do about se não houver versões
        if AboutVersion.query.count() == 0:
            print("📂 Nenhuma versão do about encontrada. Carregando dados iniciais...")
            about_data = load_about_from_json()
            if not about_data:
                about_data = get_default_about_data()
            
            initial_version = AboutVersion(
                version='1.0.0',
                data=about_data,
                is_current=True
            )
            db.session.add(initial_version)
            print("✅ Dados iniciais do about carregados!")
        
        # Carregar dados iniciais do profile se não houver versões
        if ProfileVersion.query.count() == 0:
            print("📂 Nenhuma versão do profile encontrada. Carregando dados iniciais...")
            profile_data = load_json_data('profile.json', None)
            if profile_data:
                initial_version = ProfileVersion(
                    version='1.0.0',
                    data=profile_data,
                    is_current=True
                )
                db.session.add(initial_version)
                print("✅ Dados iniciais do profile carregados!")
        
        db.session.commit()
    
    print("🚀 Servidor rodando em http://localhost:5000")
    print("📋 Endpoints disponíveis:")
    print("   GET  /api/curriculum     - Obter currículo")
    print("   PUT  /api/curriculum     - Atualizar currículo (auth)")
    print("   GET  /api/about          - Obter dados do about")
    print("   PUT  /api/about          - Atualizar dados do about (auth)")
    print("   GET  /api/profile        - Obter dados do perfil")
    print("   POST /api/auth/login     - Login")
    print("   POST /api/auth/register  - Registrar usuário")
    print("   GET  /api/auth/verify    - Verificar token (auth)")
    print("   GET  /api/health         - Status do servidor")
    print("   GET  /api/download/{ext} - Download de arquivos")
    
    app.run(debug=True, host='0.0.0.0', port=5000)
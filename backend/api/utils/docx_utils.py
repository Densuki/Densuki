# utils/docx_utils.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(data):
    """Gera um documento DOCX a partir dos dados do currículo"""
    doc = Document()
    
    # Configurar margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Título
    title = doc.add_heading('JOÃO GABRIEL SOUSA SANTOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contato
    if data.get('contact'):
        contact_texts = []
        for key, value in data['contact'].items():
            if value:
                contact_texts.append(str(value))
        if contact_texts:
            contact = doc.add_paragraph(' • '.join(contact_texts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Objetivo
    if data.get('objective'):
        doc.add_heading('OBJETIVO PROFISSIONAL', 2)
        doc.add_paragraph(data['objective'])
    
    # Resumo
    if data.get('summary'):
        doc.add_heading('RESUMO PROFISSIONAL', 2)
        doc.add_paragraph(data['summary'])
    
    # Formação Acadêmica
    if data.get('education'):
        doc.add_heading('FORMAÇÃO ACADÊMICA', 2)
        for edu in data['education']:
            text = f"• {edu.get('degree', '')}"
            if edu.get('institution'):
                text += f" – {edu.get('institution')}"
            if edu.get('period'):
                text += f" ({edu.get('period')})"
            if edu.get('status'):
                status_labels = {
                    'concluido': 'Concluído',
                    'completo': 'Concluído',
                    'incompleto': 'Em andamento',
                    'trancado': 'Trancado',
                    'interrompido': 'Interrompido'
                }
                text += f" - {status_labels.get(edu['status'], edu['status'])}"
            doc.add_paragraph(text)
    
    # Cursos e Certificações
    if data.get('courses'):
        doc.add_heading('CURSOS E CERTIFICAÇÕES', 2)
        for course in data['courses']:
            text = f"• {course.get('title', '')}"
            if course.get('duration'):
                text += f" ({course.get('duration')})"
            if course.get('institution'):
                text += f" – {course.get('institution')}"
            if course.get('status'):
                status_labels = {
                    'concluido': '✅ Concluído',
                    'andamento': '🔄 Em andamento',
                    'planejado': '📋 Planejado',
                    'cancelado': '❌ Cancelado',
                    'pendente': '⏳ Pendente'
                }
                text += f" - {status_labels.get(course['status'], course['status'])}"
            doc.add_paragraph(text)
            
            if course.get('description'):
                # Processar markdown simples para descrição
                desc = course['description']
                desc = desc.replace('**', '').replace('*', '')
                desc = desc.replace('# ', '').replace('## ', '').replace('### ', '')
                lines = desc.split('\n')
                for line in lines:
                    if line.strip().startswith('-') or line.strip().startswith('*'):
                        doc.add_paragraph(f"   {line.strip()}", style='Normal')
                    elif line.strip():
                        doc.add_paragraph(line.strip(), style='Normal')
    
    # Competências
    if data.get('skills'):
        doc.add_heading('COMPETÊNCIAS', 2)
        skills_text = []
        category_labels = {
            'core': 'Soft Skills',
            'technical': 'Hard Skills'
        }
        for category, items in data['skills'].items():
            if items:
                label = category_labels.get(category, category.capitalize())
                skills_text.append(f"{label}: {', '.join(items)}")
        doc.add_paragraph(' • '.join(skills_text))
    
    # Experiências Profissionais
    if data.get('experience'):
        doc.add_heading('EXPERIÊNCIAS PROFISSIONAIS', 2)
        for exp in data['experience']:
            p = doc.add_paragraph()
            p.add_run("Empresa: ").bold = True
            p.add_run(exp.get('company', ''))
            
            p = doc.add_paragraph()
            p.add_run("Cargo: ").bold = True
            p.add_run(exp.get('position', ''))
            
            if exp.get('location'):
                p = doc.add_paragraph()
                p.add_run("Localização: ").bold = True
                p.add_run(exp.get('location', ''))
            
            p = doc.add_paragraph()
            p.add_run("Período: ").bold = True
            p.add_run(f"{exp.get('startDate', '')} - {exp.get('endDate', '')}")
            
            for resp in exp.get('responsibilities', []):
                doc.add_paragraph(f"• {resp}", style='Normal')
            
            doc.add_paragraph()
    
    # Idiomas
    if data.get('languages'):
        doc.add_heading('IDIOMAS', 2)
        for lang in data['languages']:
            doc.add_paragraph(f"{lang.get('language', '')}: {lang.get('proficiency', '')}")
    
    return doc
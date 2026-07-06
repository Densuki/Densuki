"""
=========================================================
DOCX Builder - Layout
---------------------------------------------------------
Objetivo:
    Configurar toda a estrutura física do documento.

Responsabilidades:

    • Configurar tamanho da página
    • Configurar margens
    • Configurar cabeçalho
    • Configurar rodapé
    • Configurar propriedades da seção
    • Configurar metadados

Este módulo NÃO escreve conteúdo.

=========================================================
"""
from docx.shared import Cm
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .constants import *


# =========================================================
# CONFIGURAÇÃO PRINCIPAL
# =========================================================

def configure_layout(document):
    """
    Configura completamente o layout do documento.

    Esta deve ser a primeira função chamada
    após criar o Document().
    """

    configure_metadata(document)

    for section in document.sections:
        configure_page(section)
        configure_header(section)
        configure_footer(section)


# =========================================================
# METADADOS
# =========================================================

def configure_metadata(document):
    """
    Define os metadados internos do DOCX.
    """

    props = document.core_properties

    props.title = DOCUMENT_TITLE
    props.author = AUTHOR
    props.subject = SUBJECT
    props.category = CATEGORY
    props.language = LANGUAGE


# =========================================================
# PÁGINA
# =========================================================

def configure_page(section):
    """
    Configura tamanho da folha e margens.
    """

    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT

    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM

    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    section.gutter = Cm(0)


# =========================================================
# CABEÇALHO
# =========================================================

def configure_header(section):
    """
    Configura o cabeçalho.

    Atualmente permanece vazio.
    """

    if not ENABLE_HEADER:
        return

    header = section.header
    
    # Garante que o cabeçalho tenha pelo menos um parágrafo
    if not header.paragraphs:
        header.add_paragraph()
    
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = ""


# =========================================================
# RODAPÉ
# =========================================================

def configure_footer(section):
    """
    Configura o rodapé.

    Atualmente permanece vazio.
    """

    if not ENABLE_FOOTER:
        return

    footer = section.footer
    
    # Garante que o rodapé tenha pelo menos um parágrafo
    if not footer.paragraphs:
        footer.add_paragraph()
    
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.text = ""


# =========================================================
# NOVA SEÇÃO
# =========================================================

def create_new_section(document):
    """
    Cria uma nova seção mantendo
    exatamente o mesmo layout.
    """

    section = document.add_section(WD_SECTION.NEW_PAGE)

    configure_page(section)
    configure_header(section)
    configure_footer(section)

    return section


# =========================================================
# QUEBRA DE PÁGINA
# =========================================================

def add_page_break(document):
    """
    Adiciona uma quebra de página.
    """

    document.add_page_break()


# =========================================================
# ALTERAR MARGENS
# =========================================================

def set_margins(
    section,
    top=None,
    bottom=None,
    left=None,
    right=None
):
    """
    Permite alterar margens individualmente.
    """

    if top is not None:
        section.top_margin = top

    if bottom is not None:
        section.bottom_margin = bottom

    if left is not None:
        section.left_margin = left

    if right is not None:
        section.right_margin = right


# =========================================================
# ALTERAR TAMANHO DA PÁGINA
# =========================================================

def set_page_size(
    section,
    width,
    height
):
    """
    Permite utilizar outros formatos
    além do A4.
    """

    section.page_width = width
    section.page_height = height


# =========================================================
# ORIENTAÇÃO
# =========================================================

def set_landscape(section):
    """
    Coloca a página em modo paisagem.
    """

    width = section.page_width
    height = section.page_height

    section.page_width = height
    section.page_height = width


def set_portrait(section):
    """
    Coloca a página em modo retrato.
    """

    width = section.page_width
    height = section.page_height

    if width > height:
        section.page_width = height
        section.page_height = width
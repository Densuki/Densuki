"""
=========================================================
DOCX Builder - Builder (Parte 1/3)
---------------------------------------------------------

builder.py

Classe principal responsável por construir o documento.

=========================================================
"""

from io import BytesIO

from docx import Document

from .layout import configure_layout
from .styles import create_styles
from .constants import *

from .styles import (
    apply_paragraph_style,
    apply_run_style
)


class DocumentBuilder:

    """
    Classe responsável por construir documentos DOCX.

    Exemplo:

        builder = DocumentBuilder()

        builder.title("João Gabriel")

        builder.section("Resumo")

        builder.body("Texto...")

        buffer = builder.bytes()

    """

    # =====================================================
    # CONSTRUTOR
    # =====================================================

    def __init__(self):

        self.document = Document()

        configure_layout(self.document)

        create_styles(self.document)

    # =====================================================
    # DOCUMENTO
    # =====================================================

    @property
    def doc(self):

        return self.document

    # =====================================================
    # EXPORTAÇÃO
    # =====================================================

    def save(self, path):

        self.document.save(path)

        return self

    def bytes(self):

        buffer = BytesIO()

        self.document.save(buffer)

        buffer.seek(0)

        return buffer

    # =====================================================
    # PARÁGRAFO BASE
    # =====================================================

    def paragraph(
        self,
        text="",
        style=STYLE_BODY
    ):

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            style
        )

        if text:

            run = paragraph.add_run(str(text))

            apply_run_style(
                run,
                style
            )

        return paragraph

    # =====================================================
    # TEXTO LIVRE
    # =====================================================

    def body(
        self,
        text=""
    ):

        return self.paragraph(
            text,
            STYLE_BODY
        )

    # =====================================================
    # TÍTULO PRINCIPAL
    # =====================================================

    def title(
        self,
        text
    ):

        return self.paragraph(
            text,
            STYLE_TITLE
        )

    # =====================================================
    # CONTATO
    # =====================================================

    def contact(
        self,
        text
    ):

        return self.paragraph(
            text,
            STYLE_CONTACT
        )

    # =====================================================
    # TÍTULO DE SEÇÃO
    # =====================================================

    def section(
        self,
        text
    ):

        return self.paragraph(
            text,
            STYLE_SECTION
        )

    # =====================================================
    # ESPAÇAMENTO
    # =====================================================

    def spacing(
        self,
        size=SPACE_NORMAL
    ):

        paragraph = self.document.add_paragraph()

        paragraph.paragraph_format.space_after = size

        return paragraph

    # =====================================================
    # QUEBRA DE PÁGINA
    # =====================================================

    def page_break(self):

        self.document.add_page_break()

        return self

    # =====================================================
    # LINHA SEPARADORA
    # =====================================================

    def separator(
        self,
        character="-",
        amount=70
    ):

        return self.body(
            character * amount
        )

    # =====================================================
    # LABEL + VALUE
    # =====================================================

    def label_value(
        self,
        label,
        value,
        separator=": "
    ):
        """
        Exemplo:

            Empresa: Microsoft
        """

        paragraph = self.document.add_paragraph()

        run = paragraph.add_run(f"{label}{separator}")
        apply_run_style(run, STYLE_LABEL)

        run = paragraph.add_run(str(value))
        apply_run_style(run, STYLE_VALUE)

        return paragraph

    # =====================================================
    # EMPRESA
    # =====================================================

    def company(
        self,
        name
    ):

        return self.paragraph(
            name,
            STYLE_COMPANY
        )

    # =====================================================
    # CARGO
    # =====================================================

    def position(
        self,
        position
    ):

        return self.paragraph(
            position,
            STYLE_POSITION
        )

    # =====================================================
    # PERÍODO
    # =====================================================

    def period(
        self,
        start,
        end=None
    ):

        if not end:
            end = "Atual"

        text = f"{start} {LINE} {end}"

        return self.paragraph(
            text,
            STYLE_PERIOD
        )

    # =====================================================
    # BULLET
    # =====================================================

    def bullet(
        self,
        text
    ):

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            STYLE_BULLET
        )

        run = paragraph.add_run(
            f"{BULLET} {text}"
        )

        apply_run_style(
            run,
            STYLE_BULLET
        )

        return paragraph

    # =====================================================
    # LISTA
    # =====================================================

    def bullets(
        self,
        items
    ):

        if not items:
            return self

        for item in items:
            self.bullet(item)

        return self

    # =====================================================
    # TEXTO EM NEGRITO
    # =====================================================

    def bold(
        self,
        text,
        style=STYLE_BODY
    ):

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            style
        )

        run = paragraph.add_run(str(text))

        apply_run_style(
            run,
            style
        )

        run.bold = True

        return paragraph

    # =====================================================
    # TEXTO EM ITÁLICO
    # =====================================================

    def italic(
        self,
        text,
        style=STYLE_BODY
    ):

        paragraph = self.document.add_paragraph()

        apply_paragraph_style(
            paragraph,
            style
        )

        run = paragraph.add_run(str(text))

        apply_run_style(
            run,
            style
        )

        run.italic = True

        return paragraph

    # =====================================================
    # LINHA DE CONTATO
    # =====================================================

    def contact_line(
        self,
        values,
        separator=SEPARATOR
    ):

        items = []

        for value in values:

            if value:

                items.append(str(value))

        return self.contact(
            separator.join(items)
        )

    # =====================================================
    # TEXTO PERSONALIZADO
    # =====================================================

    def text(
        self,
        text,
        style=STYLE_BODY
    ):

        return self.paragraph(
            text,
            style
        )

    # =====================================================
    # PARÁGRAFO VAZIO
    # =====================================================

    def empty_line(self):

        self.document.add_paragraph()

        return self
   
    # =====================================================
    # UTILITÁRIOS
    # =====================================================

    @staticmethod
    def is_empty(value):
        """
        Retorna True caso o valor seja None
        ou contenha apenas espaços.
        """

        if value is None:
            return True

        return str(value).strip() == ""

    @staticmethod
    def default(value, fallback=""):

        if DocumentBuilder.is_empty(value):
            return fallback

        return value

    @staticmethod
    def clean_markdown(text):
        """
        Remove marcações básicas de Markdown.
        """

        if not text:
            return ""

        replacements = [
            ("**", ""),
            ("__", ""),
            ("#", ""),
            ("*", ""),
            ("`", ""),
            (">", ""),
        ]

        for old, new in replacements:
            text = text.replace(old, new)

        return text.strip()

    @staticmethod
    def join(values, separator=SEPARATOR):
        """
        Junta apenas valores preenchidos.
        """

        result = []

        for value in values:

            if not DocumentBuilder.is_empty(value):

                result.append(str(value))

        return separator.join(result)

    @staticmethod
    def format_period(start=None, end=None):

        start = DocumentBuilder.default(start)

        end = DocumentBuilder.default(end, "Atual")

        return f"{start} {LINE} {end}"

    # =====================================================
    # BLOCOS
    # =====================================================

    def experience(
        self,
        company,
        position,
        start,
        end=None,
        location=None,
        responsibilities=None
    ):
        """
        Cria um bloco completo de experiência.
        """

        self.company(company)
        self.position(position)
        self.period(start, end)

        if location:
            self.body(location)

        if responsibilities:
            self.bullets(responsibilities)

        self.spacing()

        return self

    def education(
        self,
        degree,
        institution="",
        period="",
        status=""
    ):

        values = [degree]

        if institution:
            values.append(institution)

        if period:
            values.append(period)

        if status:
            values.append(status)

        self.bullet(" — ".join(values))

        return self

    def course(
        self,
        title,
        institution="",
        duration="",
        status="",
        description=""
    ):

        values = [title]

        if duration:
            values.append(duration)

        if institution:
            values.append(institution)

        if status:
            values.append(status)

        self.bullet(" — ".join(values))

        if description:

            description = self.clean_markdown(description)

            for line in description.splitlines():

                line = line.strip()

                if line:

                    self.body(line)

        return self

    def language(
        self,
        language,
        proficiency
    ):

        self.bullet(
            f"{language}: {proficiency}"
        )

        return self

    # =====================================================
    # FINALIZAÇÃO
    # =====================================================

    def build(self):
        """
        Retorna o Document.
        """

        return self.document

    def reset(self):
        """
        Reinicia completamente o documento.
        """

        self.document = Document()

        configure_layout(self.document)

        create_styles(self.document)

        return self

    def __enter__(self):

        return self

    def __exit__(
        self,
        exc_type,
        exc_val,
        exc_tb
    ):

        return False